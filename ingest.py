import os
from docx import Document
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LangChainDocument

# Configuration
DATA_PATH = "data"
CHROMA_PATH = "chroma_db_local"
EMBEDDING_MODEL = "nomic-embed-text"

def load_docx(file_path):
    doc = Document(file_path)
    full_text = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract from tables (the user's request for complex structures)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    text_content = "\n".join(full_text)
    return [LangChainDocument(page_content=text_content, metadata={"source": file_path})]

def main():
    # 1. Load documents
    print(f"Loading documents from {DATA_PATH}...")
    
    documents = []
    for file in os.listdir(DATA_PATH):
        file_path = os.path.join(DATA_PATH, file)
        ext = os.path.splitext(file)[1].lower()
        
        try:
            if ext == ".txt":
                loader = TextLoader(file_path, encoding="utf-8")
                documents.extend(loader.load())
            elif ext == ".pdf":
                loader = PyPDFLoader(file_path)
                documents.extend(loader.load())
            elif ext == ".docx":
                documents.extend(load_docx(file_path))
        except Exception as e:
            print(f"Error loading {file}: {e}")
    
    if not documents:
        print("No documents found in data/ folder.")
        return

    # 2. Split text
    print("Splitting documents into chunks...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        add_start_index=True
    )
    chunks = text_splitter.split_documents(documents)
    print(f"Created {len(chunks)} chunks.")

    # 3. Create embeddings and store in Chroma
    print(f"Creating embeddings and saving to {CHROMA_PATH}...")
    embeddings = OllamaEmbeddings(model=EMBEDDING_MODEL)
    
    # Overwrite existing database
    db = Chroma.from_documents(
        documents=chunks, 
        embedding=embeddings, 
        persist_directory=CHROMA_PATH
    )
    print("Ingestion complete! You can now run the app.")

if __name__ == "__main__":
    main()
